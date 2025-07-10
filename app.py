import os
import tempfile
from flask import Flask, request, send_file, jsonify
from PIL import Image
from io import BytesIO
from docx import Document
from docx.shared import Mm
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from werkzeug.utils import secure_filename
from flask_cors import CORS
from pdf2image import convert_from_bytes
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

app = Flask(__name__)
CORS(app, origins=["https://timur-kalabin.github.io"])
app.config['MAX_CONTENT_LENGTH'] = 32 * 1024 * 1024
app.config['UPLOAD_FOLDER'] = 'uploads'

os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

def set_cell_margins(cell, top=50, bottom=50):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin, value in [('top', top), ('bottom', bottom)]:
        node = OxmlElement(f'w:{margin}')
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def process_image(img_stream):
    try:
        with Image.open(img_stream) as img:
            rotated = img.rotate(90, expand=True)
            buffer = BytesIO()
            rotated.save(buffer, format="PNG")
            buffer.seek(0)
            return buffer
    except Exception as e:
        print(f"Ошибка обработки изображения: {str(e)}")
        return None

def create_document_from_images(image_streams):
    PAGE_WIDTH = Mm(210)
    PAGE_HEIGHT = Mm(297)
    CELL_WIDTH = Mm(92)
    CELL_HEIGHT = Mm(131)
    MARGIN = Mm(5)
    TOP_SPACE = Mm(10)
    doc = Document()
    section = doc.sections[0]
    section.page_width = PAGE_WIDTH
    section.page_height = PAGE_HEIGHT
    section.left_margin = MARGIN
    section.right_margin = MARGIN
    section.top_margin = MARGIN
    section.bottom_margin = MARGIN
    for i in range(0, len(image_streams), 4):
        table = doc.add_table(rows=2, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        for row in table.rows:
            row.height = Mm(CELL_HEIGHT.mm + TOP_SPACE.mm)
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            for cell in row.cells:
                cell.width = CELL_WIDTH
                cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
                set_cell_margins(cell, top=0, bottom=0)
        for idx in range(4):
            if i + idx >= len(image_streams):
                break
            row = idx // 2
            col = idx % 2
            cell = table.cell(row, col)
            img_data = process_image(image_streams[i + idx])
            if img_data:
                para = cell.paragraphs[0]
                para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = para.add_run()
                run.add_picture(img_data, width=CELL_WIDTH, height=CELL_HEIGHT)
        if i + 4 < len(image_streams):
            doc.add_page_break()
    return doc

@app.route('/api/upload', methods=['POST'])
def upload_files():
    try:
        if 'files' not in request.files:
            return jsonify({'error': 'Файлы не выбраны'}), 400
        files = request.files.getlist('files')
        if not files or files[0].filename == '':
            return jsonify({'error': 'Файлы не выбраны'}), 400
        image_streams = []
        allowed_image_ext = {'.png', '.jpg', '.jpeg', '.gif', '.bmp'}
        for file in files:
            if file and file.filename:
                filename = secure_filename(file.filename)
                file_ext = os.path.splitext(filename)[1].lower()
                file_bytes = file.read()
                if file_ext in allowed_image_ext:
                    image_streams.append(BytesIO(file_bytes))
                elif file_ext == '.pdf':
                    # Конвертируем PDF в изображения
                    pdf_images = convert_from_bytes(file_bytes, dpi=200)
                    for img in pdf_images:
                        buf = BytesIO()
                        img.save(buf, format='PNG')
                        buf.seek(0)
                        image_streams.append(buf)
        if not image_streams:
            return jsonify({'error': 'Нет подходящих изображений или PDF'}), 400
        doc = create_document_from_images(image_streams)
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_file.name)
        temp_file.close()
        return send_file(
            temp_file.name,
            as_attachment=True,
            download_name='Images.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        return jsonify({'error': f'Ошибка при обработке: {str(e)}'}), 500

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5001) 